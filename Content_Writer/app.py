import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, AIMessage
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO  

# Load environment variables
load_dotenv()
os.environ["GROQ_API_KEY"] = os.getenv("GROQ_API_KEY")

# Initialize LLM
llm = ChatGroq(model="llama3-70b-8192")

# Function to generate multiple subtopics
def generate_subtopics(topic):
    response = llm.invoke(f"Generate as many detailed subtopics as possible on: {topic}")
    subtopics = [sub.strip() for sub in response.content.split("\n") if sub.strip() and not sub.strip().endswith(":")]
    return subtopics

# Function to generate content based on selected subtopic
def generate_content(selected_subtopic):
    response = llm.invoke(f"Write a detailed, well-structured explanation on: {selected_subtopic}")
    return response.content

# Function to create a .docx file
def create_docx(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to handle user queries related to the entire topic
def ask_question(topic, query):
    response = llm.invoke(f"Based on the topic '{topic}', answer this question: {query}")
    return response.content

# Streamlit Layout
st.set_page_config(layout="wide")

st.markdown("<h1 style='text-align: center;'>📚 AI Content Generator</h1>", unsafe_allow_html=True)

# Input for topic generation
topic = st.text_input("Enter a broad topic:")

if st.button("Generate Subtopics"):
    st.session_state["subtopics"] = generate_subtopics(topic)

st.markdown("---")

# ---- Three Column Layout ----
col1, col2, col3 = st.columns(3)

# Column 1: Subtopics Section
with col1:
    st.subheader("📌 Index")
    if "subtopics" in st.session_state:
        selected_subtopic = st.radio("Select a subtopic:", st.session_state["subtopics"])
        st.session_state["selected_subtopic"] = selected_subtopic

# Column 2: Content Display Section
with col2:
    if "selected_subtopic" in st.session_state:
        st.subheader(f"📖 {st.session_state['selected_subtopic']}")

        if st.button("Generate Content"):
            st.session_state["content"] = generate_content(st.session_state["selected_subtopic"])

        if "content" in st.session_state:
            st.markdown(st.session_state["content"])

            # Download as .docx
            docx_buffer = create_docx(st.session_state["selected_subtopic"], st.session_state["content"])
            st.download_button(
                label="📥 Download as .docx",
                data=docx_buffer,
                file_name=f"{st.session_state['selected_subtopic'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Column 3: Q&A Chat Section
with col3:
    st.subheader(f"How may I help you?")

    query = st.text_input("Ask anything related to this topic:")
    
    if st.button("Get Answer"):
        if query:
            answer = ask_question(topic, query)
            st.markdown(f"**Answer:**\n\n{answer}")
        else:
            st.warning("Please enter a question.")